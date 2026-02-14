from docx import Document

docx_path = "data/calendar_2024.docx"
doc = Document(docx_path)

print(f"Document has {len(doc.paragraphs)} paragraphs")
print(f"Document has {len(doc.tables)} tables\n")

print("=== PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs[:10]):
    print(f"{i}: {repr(para.text[:100])}")

print("\n=== TABLES ===")
for t_idx, table in enumerate(doc.tables[:3]):
    print(f"Table {t_idx}: {len(table.rows)} rows x {len(table.columns)} cols")
    for r_idx, row in enumerate(table.rows[:3]):
        for c_idx, cell in enumerate(row.cells[:3]):
            print(f"  [{r_idx},{c_idx}]: {repr(cell.text[:50])}")