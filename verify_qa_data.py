#!/usr/bin/env python3
"""Verify QA dataset consistency between documents and JSON files"""
import json
from docx import Document
from pathlib import Path

def extract_text_from_docx(path):
    """Extract text from DOCX file"""
    doc = Document(path)
    full_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        full_text.append(text)
    return "\n".join(full_text)

def verify_qa_pair(context, qa_item):
    """Check if answer_start and answer_text are correct"""
    answer_start = qa_item["answer_start"]
    answer_text = qa_item["answer_text"]
    
    # Get expected answer from context
    expected = context[answer_start:answer_start + len(answer_text)]
    
    # For multi-line answers, might need to normalize
    expected_normalized = expected.replace('\n', ' ')
    answer_normalized = answer_text.replace('\n', ' ')
    
    if expected == answer_text:
        return True, "✓ Exact match"
    elif expected_normalized.startswith(answer_normalized.split()[0]):
        return True, "✓ Partial match (first word)"
    else:
        return False, f"✗ Mismatch: expected '{expected[:30]}...' but got '{answer_text[:30]}...'"

def main():
    data_dir = Path("data")
    
    for json_file in sorted(data_dir.glob("*.json")):
        docx_file = json_file.with_suffix(".docx")
        if not docx_file.exists():
            print(f"SKIP: {json_file.name} - no matching DOCX")
            continue
        
        print(f"\n{'='*60}")
        print(f"Checking: {json_file.name}")
        print(f"{'='*60}")
        
        try:
            # Extract context
            context = extract_text_from_docx(docx_file)
            print(f"Document text length: {len(context)} chars")
            
            # Load QA pairs
            with open(json_file) as f:
                qa_items = json.load(f)
            print(f"QA pairs: {len(qa_items)}")
            
            # Verify each pair
            valid_count = 0
            invalid_count = 0
            
            for i, qa in enumerate(qa_items[:5]):  # Check first 5
                is_valid, msg = verify_qa_pair(context, qa)
                if is_valid:
                    valid_count += 1
                    print(f"  [{i}] {msg}")
                else:
                    invalid_count += 1
                    print(f"  [{i}] {msg}")
                    print(f"       Q: {qa['question'][:50]}...")
                    print(f"       A: {qa['answer_text'][:50]}...")
            
            print(f"\nSummary: {valid_count} valid, {invalid_count} invalid (out of 5 checked)")
            
        except Exception as e:
            print(f"ERROR: {e}")

if __name__ == "__main__":
    main()
